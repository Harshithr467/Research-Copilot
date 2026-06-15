import os
import tempfile

import docx
import fitz
import pytest

from extraction.extractor import ExtractionError, extract


# ── helpers ───────────────────────────────────────────────────────────────────

def _make_pdf(path: str, pages: list[str]) -> None:
    """Create a minimal PDF; empty string → blank page (simulates scanned)."""
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        if text:
            page.insert_text((50, 50), text)
    doc.save(path)
    doc.close()


def _make_docx(path: str, paragraphs: list[str], table_rows: list[list[str]]) -> None:
    """Create a DOCX with the given paragraphs and a single table."""
    doc = docx.Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        cols = max(len(row) for row in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        for r, row in enumerate(table_rows):
            for c, cell_text in enumerate(row):
                table.cell(r, c).text = cell_text
    doc.save(path)


# ── PDF ───────────────────────────────────────────────────────────────────────

def test_pdf_one_record_per_page():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "paper.pdf")
        _make_pdf(path, ["Introduction text.", "Methods text."])

        chunks = extract(path)

        assert len(chunks) == 2
        assert chunks[0]["page"] == 1
        assert chunks[1]["page"] == 2
        assert "Introduction" in chunks[0]["text"]
        assert chunks[0]["source"] == "paper.pdf"
        assert chunks[1]["source"] == "paper.pdf"


def test_pdf_skips_blank_pages_keeps_text_pages(caplog):
    import logging
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "mixed.pdf")
        _make_pdf(path, ["Real content.", "", "More content."])  # page 2 is blank

        with caplog.at_level(logging.WARNING, logger="extraction.extractor"):
            chunks = extract(path)

        assert len(chunks) == 2
        assert chunks[0]["page"] == 1
        assert chunks[1]["page"] == 3  # page 2 was skipped
        assert "skipped" in caplog.text
        assert "[2]" in caplog.text


def test_pdf_all_blank_raises():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "scanned.pdf")
        _make_pdf(path, ["", ""])

        with pytest.raises(ExtractionError, match="image-only"):
            extract(path)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def test_docx_paragraphs_and_table_cells_in_single_record():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "paper.docx")
        _make_docx(
            path,
            paragraphs=["Abstract paragraph.", "Introduction paragraph."],
            table_rows=[
                ["Method", "Accuracy", "F1"],
                ["Baseline", "0.82", "0.79"],
                ["Ours", "0.91", "0.89"],
            ],
        )

        chunks = extract(path)

        assert len(chunks) == 1
        chunk = chunks[0]
        assert chunk["page"] is None
        assert chunk["source"] == "paper.docx"
        # paragraph text must be present
        assert "Abstract paragraph." in chunk["text"]
        assert "Introduction paragraph." in chunk["text"]
        # table cell text must also be present
        assert "Method" in chunk["text"]
        assert "Baseline" in chunk["text"]
        assert "0.91" in chunk["text"]


# ── TXT / MD ──────────────────────────────────────────────────────────────────

def test_txt_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "notes.txt")
        with open(path, "w", encoding="utf-8") as f:
            f.write("Hello world\nLine two")

        chunks = extract(path)

        assert len(chunks) == 1
        assert chunks[0]["page"] is None
        assert "Hello world" in chunks[0]["text"]
        assert chunks[0]["source"] == "notes.txt"


def test_md_extraction():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "readme.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write("# Title\n\nSome content here.")

        chunks = extract(path)

        assert len(chunks) == 1
        assert "Title" in chunks[0]["text"]
        assert chunks[0]["page"] is None


# ── error cases ───────────────────────────────────────────────────────────────

def test_empty_txt_raises():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "empty.txt")
        open(path, "w").close()

        with pytest.raises(ExtractionError):
            extract(path)


def test_garbage_pdf_raises():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "garbage.pdf")
        with open(path, "wb") as f:
            f.write(b"\x00\x01\x02\x03 not a real pdf")

        with pytest.raises(ExtractionError):
            extract(path)


def test_unsupported_extension_raises():
    with tempfile.TemporaryDirectory() as tmp:
        path = os.path.join(tmp, "data.csv")
        with open(path, "w") as f:
            f.write("a,b,c")

        with pytest.raises(ExtractionError, match="Unsupported file type"):
            extract(path)


def test_missing_file_raises():
    with pytest.raises(ExtractionError, match="File not found"):
        extract("/tmp/does_not_exist_xyz.pdf")
