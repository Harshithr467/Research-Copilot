import logging
import os
from typing import TypedDict

import fitz   # PyMuPDF
import docx   # python-docx

logger = logging.getLogger(__name__)


class ExtractedChunk(TypedDict):
    text: str
    source: str       # basename of the originating file
    page: int | None  # 1-indexed page number (PDFs only); None for all other formats


class ExtractionError(Exception):
    """Raised when a document cannot be extracted: corrupt, unreadable, or fully image-only."""


def extract(path: str) -> list[ExtractedChunk]:
    """
    Detect document type by extension and extract text with source metadata.

    Returns one ExtractedChunk per page for PDFs; one record for DOCX/TXT/MD.
    Raises ExtractionError for unsupported types, unreadable files, or documents
    that yield zero extractable text across all pages.
    """
    if not os.path.isfile(path):
        raise ExtractionError(f"File not found: {path!r}")

    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    if ext in (".txt", ".md"):
        return _extract_txt(path)
    raise ExtractionError(f"Unsupported file type: {ext!r}")


def _extract_pdf(path: str) -> list[ExtractedChunk]:
    source = os.path.basename(path)
    try:
        doc = fitz.open(path)
    except Exception as e:
        raise ExtractionError(f"Cannot open PDF {path!r}: {e}") from e

    records: list[ExtractedChunk] = []
    skipped_pages: list[int] = []

    for i in range(len(doc)):
        text = doc[i].get_text()
        if text.strip():
            records.append({"text": text, "source": source, "page": i + 1})
        else:
            skipped_pages.append(i + 1)

    doc.close()

    if not records:
        raise ExtractionError(
            f"No extractable text in {path!r}. "
            "All pages are empty or image-only (scanned PDF with no text layer)."
        )

    if skipped_pages:
        logger.warning(
            "%s: skipped %d empty/image-only page(s): %s",
            source,
            len(skipped_pages),
            skipped_pages,
        )

    return records


def _extract_docx(path: str) -> list[ExtractedChunk]:
    source = os.path.basename(path)
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise ExtractionError(f"Cannot open DOCX {path!r}: {e}") from e

    parts: list[str] = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # python-docx paragraph iteration skips table content; extract it explicitly.
    # Merged cells appear multiple times in row.cells — deduplicate by XML element identity.
    # Use `is` comparison on stored references rather than id(): id() on a temporary
    # lxml proxy can be reused after GC, giving false duplicate matches.
    seen: set = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc = cell._tc  # hold strong reference so the proxy can't be GC'd and its address reused
                if tc not in seen:
                    seen.add(tc)
                    if cell.text.strip():
                        parts.append(cell.text)

    if not parts:
        raise ExtractionError(f"No extractable text found in {path!r}.")

    return [{"text": "\n".join(parts), "source": source, "page": None}]


def _extract_txt(path: str) -> list[ExtractedChunk]:
    source = os.path.basename(path)
    try:
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()
    except Exception as e:
        raise ExtractionError(f"Cannot read {path!r}: {e}") from e

    if not text.strip():
        raise ExtractionError(f"File {path!r} is empty or contains no readable text.")

    return [{"text": text, "source": source, "page": None}]
